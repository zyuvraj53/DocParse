import pdfplumber
import docx
import pytesseract
from PIL import Image
import re
import json
import os
from pathlib import Path

def extract_text_from_pdf(file_path):
    """Extract text from PDF files using pdfplumber."""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return None

def extract_text_from_docx(file_path):
    """Extract text from DOCX files using python-docx."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += "\n" + " | ".join(row_text)
        
        return text
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return None

def extract_text_from_image(file_path):
    """Extract text from images using pytesseract."""
    try:
        text = pytesseract.image_to_string(Image.open(file_path))
        return text
    except Exception as e:
        print(f"Error reading image {file_path}: {e}")
        return None

def parse_payslip(text):
    """Parse payslip text to extract components and validate."""
    if not text:
        return {"error": "No text extracted from the document"}

    # Clean up the text for better parsing
    text = text.replace('\t', ' ').replace('₹', '').replace(',', '')
    
    # Define multiple regex patterns for different formats (case-insensitive)
    patterns = {
        # Basic salary patterns - multiple variations
        "basic": [
            r"basic\s*(?:pay|salary)?\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)",
            r"basic\s*(\d{1,6}(?:\.\d{2})?)",
            r"(?:basic\s*pay|basic\s*salary)\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # HRA patterns
        "hra": [
            r"(?:hra|house\s*rent\s*allowance)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)",
            r"house\s*rent\s*allowance\s*(\d{1,6}(?:\.\d{2})?)",
            r"hra\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # Variable pay / incentives / bonuses
        "variable_pay": [
            r"(?:variable\s*pay|incentive\s*pay|bonus|other\s*allowance|i[cn]entive)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)",
            r"(?:meal\s*allowance|transport\s*allowance|special\s*allowance)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # Total earnings
        "total_earnings": [
            r"(?:total\s*earnings?|gross\s*earnings?|gross\s*pay|total\s*pay)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # Total deductions
        "total_deductions": [
            r"(?:total\s*deductions?|total\s*deduction)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # Net pay patterns
        "net_pay": [
            r"net\s*pay[|\s]*(\d{1,6}(?:\.\d{2})?)",
            r"(?:net\s*(?:salary|payable)|total\s*net\s*payable|employee\s*net\s*pay)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)",
            r"(?:take\s*home|net\s*amount)\s*[:\-|]?\s*(\d{1,6}(?:\.\d{2})?)"
        ],
        
        # Employee details
        "employee_name": [
            r"(?:employee\s*name|name)\s*[:\-]?\s*([A-Za-z][A-Za-z\s]{1,30})",
            r"name\s*:\s*([A-Za-z][A-Za-z\s]{1,30})"
        ],
        
        "employee_id": [
            r"(?:employee\s*id|emp\s*id|id)\s*[:\-]?\s*(\w+)",
            r"employee\s*id\s*:\s*(\w+)"
        ],
        
        "designation": [
            r"designation\s*[:\-]?\s*([A-Za-z][A-Za-z\s]{1,40})",
            r"(?:position|role|title)\s*[:\-]?\s*([A-Za-z][A-Za-z\s]{1,40})"
        ]
    }

    result = {
        "components": {},
        "employment_proof": {}
    }

    # Extract components using multiple regex patterns
    for key, pattern_list in patterns.items():
        found = False
        for pattern in pattern_list:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                if key in ["employee_name", "employee_id", "designation"]:
                    # Clean up the extracted text
                    extracted_text = matches[0].strip() if isinstance(matches[0], str) else str(matches[0]).strip()
                    # Remove common suffixes and clean up
                    extracted_text = re.sub(r'\s*(pf\s*no|employee|department|designation).*$', '', extracted_text, flags=re.IGNORECASE).strip()
                    # Remove newlines and extra content
                    if '\n' in extracted_text:
                        extracted_text = extracted_text.split('\n')[0].strip()
                    if extracted_text and len(extracted_text) > 1 and not extracted_text.isdigit():
                        result["employment_proof"][key] = extracted_text
                        found = True
                        break
                else:
                    # For numeric values, take the last (most final) match for net_pay
                    # but first match for other components
                    for match in (matches if key != "net_pay" else [matches[-1]]):
                        try:
                            value = float(str(match).replace(',', ''))
                            if value > 0:  # Only positive values
                                result["components"][key] = value
                                found = True
                                break
                        except (ValueError, TypeError):
                            continue
                if found:
                    break

    # Additional extraction for standalone numbers (like net pay at the end)
    if "net_pay" not in result["components"]:
        # Look for the last occurrence of Net Pay followed by amount
        net_pay_matches = re.findall(r'net\s*pay[|\s]*(\d{1,6}(?:\.\d{2})?)', text, re.IGNORECASE)
        if net_pay_matches:
            # Take the last (final) net pay amount
            result["components"]["net_pay"] = float(net_pay_matches[-1])
        else:
            # Look for standalone large numbers that could be net pay (at the end)
            text_lines = text.strip().split('\n')
            for line in reversed(text_lines[-5:]):  # Check last 5 lines
                standalone_numbers = re.findall(r'\b(\d{4,6}(?:\.\d{2})?)\b', line)
                if standalone_numbers:
                    potential_net_pay = float(standalone_numbers[-1])
                    if 1000 <= potential_net_pay <= 500000:
                        result["components"]["net_pay"] = potential_net_pay
                        break

    # Calculate missing values
    total_earnings = result["components"].get("total_earnings", 0)
    total_deductions = result["components"].get("total_deductions", 0)
    net_pay = result["components"].get("net_pay", 0)
    
    # If we have total earnings and deductions but no net pay, calculate it
    if total_earnings > 0 and total_deductions >= 0 and net_pay == 0:
        result["components"]["net_pay"] = total_earnings - total_deductions
        net_pay = result["components"]["net_pay"]
    
    # If we have net pay and deductions but no total earnings, calculate it
    elif net_pay > 0 and total_deductions >= 0 and total_earnings == 0:
        result["components"]["total_earnings"] = net_pay + total_deductions
    
    # Set net_salary as alias for net_pay
    result["components"]["net_salary"] = result["components"].get("net_pay", 0)

    # Validate employment proof
    result["employment_proof"]["valid"] = bool(
        result["employment_proof"].get("employee_name") or
        result["employment_proof"].get("employee_id")
    )

    return result

def process_payslip(file_path):
    """Process payslip based on file extension."""
    ext = Path(file_path).suffix.lower()
    text = None

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext in [".png", ".jpg", ".jpeg"]:
        text = extract_text_from_image(file_path)
    else:
        return {"error": f"Unsupported file format: {ext} for {file_path}"}

    result = parse_payslip(text)
    result["file_processed"] = file_path
    return result

def save_to_json(data, output_path):
    """Save parsed data to JSON file."""
    try:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w") as f:
            json.dump(data, f, indent=4)
        print(f"Output saved to {output_path}")
    except Exception as e:
        print(f"Error saving JSON for {output_path}: {e}")

def main():
    """Process all payslips in the uploads folder."""
    uploads_dir = "uploads_payslips"
    outputs_dir = "outputs_payslips"
    if not os.path.exists(uploads_dir):
        print(f"Error: {uploads_dir} folder not found")
        return

    for file in os.listdir(uploads_dir):
        file_path = os.path.join(uploads_dir, file)
        if os.path.isfile(file_path):
            output_path = os.path.join(outputs_dir, f"output_{file}.json")
            result = process_payslip(file_path)
            save_to_json(result, output_path)

if __name__ == "__main__":
    main()